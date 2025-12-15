#!/usr/bin/env python3
"""
Extract placeholders from F-103.docx template
"""

import sys
from docx import Document
import re

def extract_placeholders(docx_path):
    """Extract all {placeholder} patterns from DOCX file"""
    doc = Document(docx_path)
    placeholders = set()
    
    # Extract from paragraphs
    for para in doc.paragraphs:
        text = para.text
        matches = re.findall(r'\{([^}]+)\}', text)
        placeholders.update(matches)
    
    # Extract from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    text = para.text
                    matches = re.findall(r'\{([^}]+)\}', text)
                    placeholders.update(matches)
    
    return sorted(placeholders)

if __name__ == '__main__':
    if len(sys.argv) < 2:
        print("Usage: python extract_placeholders_f103.py <path_to_F-103.docx>")
        sys.exit(1)
    
    docx_path = sys.argv[1]
    
    print("=" * 70)
    print("PLACEHOLDERS FROM F-103.docx")
    print("=" * 70)
    print()
    
    placeholders = extract_placeholders(docx_path)
    
    print(f"Total placeholders found: {len(placeholders)}\n")
    
    for i, placeholder in enumerate(placeholders, 1):
        print(f"{i:2}. {{{placeholder}}}")
    
    print()
    print("=" * 70)
    
    # Save to file
    output_file = "documentation/F-103_PLACEHOLDERS.md"
    with open(output_file, 'w', encoding='utf-8') as f:
        f.write("# F-103.docx Template Placeholders\n\n")
        f.write("## Daftar Placeholder\n\n")
        f.write(f"Total: {len(placeholders)} placeholders\n\n")
        
        for i, placeholder in enumerate(placeholders, 1):
            f.write(f"{i}. `{{{placeholder}}}`\n")
        
        f.write("\n## Mapping ke Form Data\n\n")
        f.write("```typescript\n")
        f.write("const templateData = {\n")
        for placeholder in placeholders:
            f.write(f"  {placeholder}: formData.{placeholder} || '',\n")
        f.write("};\n")
        f.write("```\n")
    
    print(f"✅ Documentation saved to: {output_file}")
