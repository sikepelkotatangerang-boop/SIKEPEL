from docx import Document
import re

# Read the document
doc = Document('public/template/SKU.docx')

# Set to store unique placeholders
placeholders = set()

# Extract from paragraphs
for paragraph in doc.paragraphs:
    # Find all placeholders in format {placeholder_name}
    found = re.findall(r'\{([^}]+)\}', paragraph.text)
    placeholders.update(found)

# Extract from tables
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            found = re.findall(r'\{([^}]+)\}', cell.text)
            placeholders.update(found)

# Print sorted placeholders
print("=== PLACEHOLDERS IN SKU.docx ===\n")
for placeholder in sorted(placeholders):
    print(f"{{{placeholder}}}")

print(f"\n=== TOTAL: {len(placeholders)} placeholders ===")
