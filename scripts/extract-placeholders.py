#!/usr/bin/env python3
"""
Script to extract placeholders from DOCX template
Usage: python scripts/extract-placeholders.py public/template/SURATKELUAR.docx
"""

import sys
import re
from docx import Document

def extract_placeholders(docx_path):
    """Extract all {placeholder} patterns from a DOCX file"""
    try:
        doc = Document(docx_path)
        placeholders = set()
        
        # Extract from paragraphs
        for paragraph in doc.paragraphs:
            matches = re.findall(r'\{([^}]+)\}', paragraph.text)
            placeholders.update(matches)
        
        # Extract from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    matches = re.findall(r'\{([^}]+)\}', cell.text)
                    placeholders.update(matches)
        
        return sorted(list(placeholders))
    
    except Exception as e:
        print(f"Error: {e}")
        return []

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python extract-placeholders.py <path-to-docx>")
        sys.exit(1)
    
    docx_path = sys.argv[1]
    placeholders = extract_placeholders(docx_path)
    
    print(f"\n📄 Found {len(placeholders)} placeholders in {docx_path}:\n")
    for i, placeholder in enumerate(placeholders, 1):
        print(f"{i}. {{{placeholder}}}")
    
    print("\n✅ Done!")
