#!/usr/bin/env python3
"""Extract detailed content from F-103.docx"""

import re
from docx import Document

doc = Document('public/template/F-103.docx')

print("=" * 80)
print("PARAGRAPHS:")
print("=" * 80)
for i, para in enumerate(doc.paragraphs):
    if para.text.strip():
        print(f"{i}: {para.text}")

print("\n" + "=" * 80)
print("TABLES:")
print("=" * 80)
for table_idx, table in enumerate(doc.tables):
    print(f"\n--- Table {table_idx + 1} ---")
    for row_idx, row in enumerate(table.rows):
        cells_text = [cell.text.strip() for cell in row.cells]
        print(f"Row {row_idx}: {cells_text}")

print("\n" + "=" * 80)
print("PLACEHOLDERS:")
print("=" * 80)
placeholders = set()
for para in doc.paragraphs:
    matches = re.findall(r'\{([^}]+)\}', para.text)
    placeholders.update(matches)
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            matches = re.findall(r'\{([^}]+)\}', cell.text)
            placeholders.update(matches)

for p in sorted(placeholders):
    print(f"  {{{p}}}")
