#!/usr/bin/env python3
"""
Universal DOCX Template Placeholder Extractor
==============================================

Extract all {placeholder} patterns from any DOCX template file.
Generates documentation and TypeScript code snippets.

Usage:
    python extract_template_placeholders.py <path_to_template.docx>
    python extract_template_placeholders.py public/template/SKTM.docx
    python extract_template_placeholders.py public/template/F-103.docx

Features:
- Extracts placeholders from paragraphs and tables
- Handles nested tables
- Generates markdown documentation
- Creates TypeScript interface and mapping code
- Supports special characters in placeholders
- Detects duplicate placeholders
- Shows placeholder locations (paragraph/table)

Author: AI Assistant
Date: 2025-01-20
"""

import sys
import os
import re
from docx import Document
from collections import defaultdict
from pathlib import Path

class PlaceholderExtractor:
    def __init__(self, docx_path):
        self.docx_path = docx_path
        self.template_name = Path(docx_path).stem
        self.placeholders = {}  # {placeholder: [locations]}
        self.doc = None
        
    def load_document(self):
        """Load DOCX document"""
        try:
            self.doc = Document(self.docx_path)
            print(f"✅ Loaded: {self.docx_path}")
            return True
        except Exception as e:
            print(f"❌ Error loading document: {e}")
            return False
    
    def extract_from_text(self, text, location):
        """Extract placeholders from text and record location"""
        # Pattern: {anything_inside_curly_braces}
        # Supports letters, numbers, underscores, slashes, hyphens
        pattern = r'\{([a-zA-Z0-9_/\-\.]+)\}'
        matches = re.findall(pattern, text)
        
        for match in matches:
            if match not in self.placeholders:
                self.placeholders[match] = []
            self.placeholders[match].append(location)
    
    def extract_from_paragraphs(self):
        """Extract placeholders from all paragraphs"""
        for i, para in enumerate(self.doc.paragraphs):
            text = para.text
            if '{' in text and '}' in text:
                self.extract_from_text(text, f"Paragraph {i+1}")
    
    def extract_from_tables(self):
        """Extract placeholders from all tables (including nested)"""
        for table_idx, table in enumerate(self.doc.tables):
            self._extract_from_table(table, table_idx + 1)
    
    def _extract_from_table(self, table, table_num, parent_location=""):
        """Recursively extract from table and nested tables"""
        for row_idx, row in enumerate(table.rows):
            for cell_idx, cell in enumerate(row.cells):
                location = f"Table {table_num}, Row {row_idx+1}, Cell {cell_idx+1}"
                if parent_location:
                    location = f"{parent_location} > {location}"
                
                # Extract from cell paragraphs
                for para in cell.paragraphs:
                    text = para.text
                    if '{' in text and '}' in text:
                        self.extract_from_text(text, location)
                
                # Handle nested tables
                for nested_table in cell.tables:
                    self._extract_from_table(nested_table, table_num, location)
    
    def extract_all(self):
        """Extract all placeholders from document"""
        if not self.load_document():
            return False
        
        print("🔍 Extracting placeholders...")
        self.extract_from_paragraphs()
        self.extract_from_tables()
        
        print(f"✅ Found {len(self.placeholders)} unique placeholders")
        return True
    
    def get_sorted_placeholders(self):
        """Get sorted list of placeholders"""
        return sorted(self.placeholders.keys())
    
    def generate_markdown_doc(self):
        """Generate markdown documentation"""
        placeholders = self.get_sorted_placeholders()
        
        md = f"# {self.template_name}.docx - Template Placeholders\n\n"
        md += f"**Generated**: {self._get_timestamp()}\n\n"
        md += f"**Total Placeholders**: {len(placeholders)}\n\n"
        
        # Table of contents
        md += "## Table of Contents\n\n"
        md += "1. [Placeholder List](#placeholder-list)\n"
        md += "2. [Placeholder Locations](#placeholder-locations)\n"
        md += "3. [TypeScript Interface](#typescript-interface)\n"
        md += "4. [Template Data Mapping](#template-data-mapping)\n"
        md += "5. [Usage Example](#usage-example)\n\n"
        
        # Placeholder list
        md += "## Placeholder List\n\n"
        md += f"Total: **{len(placeholders)}** placeholders\n\n"
        
        for i, placeholder in enumerate(placeholders, 1):
            md += f"{i}. `{{{placeholder}}}`\n"
        
        # Placeholder locations
        md += "\n## Placeholder Locations\n\n"
        md += "Shows where each placeholder appears in the document:\n\n"
        
        for placeholder in placeholders:
            locations = self.placeholders[placeholder]
            md += f"### `{{{placeholder}}}`\n\n"
            if len(locations) > 1:
                md += f"**Appears {len(locations)} times:**\n\n"
            for loc in locations:
                md += f"- {loc}\n"
            md += "\n"
        
        # TypeScript interface
        md += "## TypeScript Interface\n\n"
        md += "```typescript\n"
        md += f"interface {self._to_pascal_case(self.template_name)}FormData {{\n"
        
        for placeholder in placeholders:
            # Convert placeholder to camelCase field name
            field_name = self._to_camel_case(placeholder)
            md += f"  {field_name}: string;\n"
        
        md += "}\n```\n\n"
        
        # Template data mapping
        md += "## Template Data Mapping\n\n"
        md += "```typescript\n"
        md += "const templateData = {\n"
        
        for placeholder in placeholders:
            # Handle special characters in placeholder names
            if '/' in placeholder or '-' in placeholder or '.' in placeholder:
                md += f"  '{placeholder}': formData.{self._to_camel_case(placeholder)} || '',\n"
            else:
                md += f"  {placeholder}: formData.{placeholder} || '',\n"
        
        md += "};\n```\n\n"
        
        # Usage example
        md += "## Usage Example\n\n"
        md += "```typescript\n"
        md += "import Docxtemplater from 'docxtemplater';\n"
        md += "import PizZip from 'pizzip';\n"
        md += "import { readFileSync } from 'fs';\n"
        md += "import { join } from 'path';\n\n"
        
        md += "// Load template\n"
        md += f"const templatePath = join(process.cwd(), 'public', 'template', '{self.template_name}.docx');\n"
        md += "const content = readFileSync(templatePath, 'binary');\n\n"
        
        md += "const zip = new PizZip(content);\n"
        md += "const doc = new Docxtemplater(zip, {\n"
        md += "  paragraphLoop: true,\n"
        md += "  linebreaks: true,\n"
        md += "  nullGetter: function() {\n"
        md += "    return '';\n"
        md += "  },\n"
        md += "});\n\n"
        
        md += "// Prepare template data\n"
        md += "const templateData = {\n"
        
        # Show first 5 placeholders as example
        for placeholder in placeholders[:5]:
            if '/' in placeholder or '-' in placeholder:
                md += f"  '{placeholder}': formData.{self._to_camel_case(placeholder)} || '',\n"
            else:
                md += f"  {placeholder}: formData.{placeholder} || '',\n"
        
        if len(placeholders) > 5:
            md += "  // ... (see Template Data Mapping section for complete list)\n"
        
        md += "};\n\n"
        
        md += "// Render document\n"
        md += "doc.render(templateData);\n\n"
        
        md += "// Generate DOCX buffer\n"
        md += "const buffer = doc.getZip().generate({\n"
        md += "  type: 'nodebuffer',\n"
        md += "  compression: 'DEFLATE',\n"
        md += "});\n"
        md += "```\n\n"
        
        # Notes
        md += "## Notes\n\n"
        md += "- All placeholders use the format `{placeholder_name}`\n"
        md += "- Empty strings (`''`) are used as default values\n"
        md += "- Special characters in placeholder names require quotes in object keys\n"
        md += f"- Template file: `public/template/{self.template_name}.docx`\n\n"
        
        # Warnings for special cases
        special_chars = [p for p in placeholders if '/' in p or '-' in p or '.' in p]
        if special_chars:
            md += "### ⚠️ Special Characters\n\n"
            md += "The following placeholders contain special characters and require quotes:\n\n"
            for p in special_chars:
                md += f"- `{{{p}}}` → `'{p}': formData.{self._to_camel_case(p)}`\n"
            md += "\n"
        
        return md
    
    def generate_console_output(self):
        """Generate formatted console output"""
        placeholders = self.get_sorted_placeholders()
        
        output = "=" * 70 + "\n"
        output += f"PLACEHOLDERS FROM {self.template_name}.docx\n"
        output += "=" * 70 + "\n\n"
        
        output += f"Total placeholders found: {len(placeholders)}\n\n"
        
        for i, placeholder in enumerate(placeholders, 1):
            output += f"{i:2}. {{{placeholder}}}\n"
        
        output += "\n" + "=" * 70 + "\n"
        
        return output
    
    def save_documentation(self, output_dir="documentation"):
        """Save markdown documentation to file"""
        os.makedirs(output_dir, exist_ok=True)
        
        filename = f"{self.template_name.upper()}_PLACEHOLDERS.md"
        filepath = os.path.join(output_dir, filename)
        
        md_content = self.generate_markdown_doc()
        
        with open(filepath, 'w', encoding='utf-8') as f:
            f.write(md_content)
        
        return filepath
    
    def _to_camel_case(self, text):
        """Convert placeholder to camelCase"""
        # Replace special characters with underscore
        text = text.replace('/', '_').replace('-', '_').replace('.', '_')
        
        # Split by underscore
        parts = text.split('_')
        
        # First part lowercase, rest capitalized
        if len(parts) == 1:
            return parts[0].lower()
        
        return parts[0].lower() + ''.join(word.capitalize() for word in parts[1:])
    
    def _to_pascal_case(self, text):
        """Convert to PascalCase"""
        text = text.replace('/', '_').replace('-', '_').replace('.', '_')
        parts = text.split('_')
        return ''.join(word.capitalize() for word in parts)
    
    def _get_timestamp(self):
        """Get current timestamp"""
        from datetime import datetime
        return datetime.now().strftime("%Y-%m-%d %H:%M:%S")


def main():
    """Main function"""
    if len(sys.argv) < 2:
        print("=" * 70)
        print("DOCX Template Placeholder Extractor")
        print("=" * 70)
        print("\nUsage:")
        print("  python extract_template_placeholders.py <path_to_template.docx>\n")
        print("Examples:")
        print("  python extract_template_placeholders.py public/template/SKTM.docx")
        print("  python extract_template_placeholders.py public/template/F-103.docx")
        print("  python extract_template_placeholders.py public\\template\\SKU.docx\n")
        sys.exit(1)
    
    docx_path = sys.argv[1]
    
    # Validate file exists
    if not os.path.exists(docx_path):
        print(f"❌ Error: File not found: {docx_path}")
        sys.exit(1)
    
    # Validate file extension
    if not docx_path.lower().endswith('.docx'):
        print(f"❌ Error: File must be a .docx file: {docx_path}")
        sys.exit(1)
    
    # Extract placeholders
    extractor = PlaceholderExtractor(docx_path)
    
    if not extractor.extract_all():
        sys.exit(1)
    
    # Print to console
    print("\n" + extractor.generate_console_output())
    
    # Save documentation
    try:
        filepath = extractor.save_documentation()
        print(f"✅ Documentation saved to: {filepath}")
        print(f"📄 Open the file to see detailed mapping and usage examples")
    except Exception as e:
        print(f"❌ Error saving documentation: {e}")
        sys.exit(1)


if __name__ == '__main__':
    main()
